#Extraction complete : glossaire (tableau) + scenarios (texte complet)

from docx import Document
from pathlib import Path


def extract_glossary_from_table(docx_path: str):
    doc = Document(docx_path)
    
    glossary_terms = []
    
    for table in doc.tables:
        first_row = table.rows[0]
        first_cell = first_row.cells[0].text.strip()
        
        if "Term" in first_cell or "Expression" in first_cell:
            
            for row in table.rows[1:]:
                term = row.cells[0].text.strip()
                if term and len(term) > 0:
                    term_clean = ' '.join(term.split())
                    glossary_terms.append(term_clean)
            
            break
    
    glossary_terms = sorted(set(glossary_terms))
    print(f"{len(glossary_terms)} termes extraits")
    return glossary_terms


def extract_all_paragraphs(docx_path: str):
    #Extrait tous les paragraphes du document (pour les scenarios)

    doc = Document(docx_path)
    
    all_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text.append(text)
    
    print(f"{len(all_text)} lignes extraites")
    return all_text


def save_glossary(terms, output_path: str):
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        for term in terms:
            f.write(f"{term}\n")
    
    print(f"Glossaire sauvegarde : {output_path}")


def save_scenarios(lines, output_path: str):
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        for line in lines:
            f.write(line + '\n')
    
    print(f"Scenarios sauvegardes : {output_path}")





if __name__ == "__main__":
    print("EXTRACTION COMPLETE DU DOCUMENT")
    
    docx_path = "data/Tomatoes-Scenarios-English-Version.docx"
    
    if not Path(docx_path).exists():
        print(f"ERREUR : Fichier non trouve : {docx_path}")
        exit(1)
    
    print(f"\nFichier trouve : {docx_path}\n")
    
    # 1. Extraction du glossaire depuis le tableau
    print("Extraction du glossaire:")
    glossary = extract_glossary_from_table(docx_path)
    save_glossary(glossary, "data/glossaires/glossaire_tomates.txt")
    
    # 2. Extraction de tous les scenarios
    print("\nExtraction des scenarios:")
    all_paragraphs = extract_all_paragraphs(docx_path)
    save_scenarios(all_paragraphs, "data/scenarios/scenarios_tomates_complet.txt")
    
    # Resultats
    print("\n")
    print("EXTRACTION TERMINEE")
    print(f"Glossaire : {len(glossary)} termes")
    print(f"Scenarios : {len(all_paragraphs)} lignes")

